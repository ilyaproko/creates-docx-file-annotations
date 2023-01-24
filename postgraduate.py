# This Python file uses the following encoding: utf-8

# ! для лучшей readalbe кода в VScode рекомендуется установить расширение 'Better Comments'

import xlrd, xlwt
import datetime
import docx
import re
import os
from docx import Document
from docx.shared import Inches, Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# * окрытие excel файла
book = xlrd.open_workbook("settings_postgraduate.xlsx")
# * первый лист excel файла
sheet = book.sheet_by_index(0)

# * создание директории для аспирантов
os.makedirs('А', exist_ok=True)

list_of_eduprogs = []

dict_of_plans = {}

# фильтровка только аспирантов, затем добавление аспирантов в массив list_of_eduprogs и добавление ключа "учебный план" в словарь
# dict_of_plans со значение 1 (нужет только чтобы сформировать запись в словаре)
for row_index in range(1, sheet.nrows): #sheet.nrows 
    # ! третья колонка в excel файле обезательно дожна представлять  ->  'учебный план' ->  пример 2.3.1-22.plx
    # * сплит наименование учебного файла с кодом, где первые ДВА элемента должны быть длинной 1, это будет указывать что это аспирантура
    educational_plan = sheet.row_values(row_index)[2].split(".");
    if len(educational_plan[0]) == 1 and len(educational_plan[1]) == 1:
        list_of_eduprogs.append(sheet.row_values(row_index)); # * добавление строки аспиранта

        # * добавление учебного плана для аспирантуры в словарь dict_of_plans
        dict_of_plans[sheet.row_values(row_index)[2]] = 1;


# * перебор только по ключам в словаре где значения представляют -> "учебный план"
for l in dict_of_plans: 
    
    # * получение всех дисциплин, практик, факультатов по ОПРЕДЕЛЕННОМУ учебному плану аспирантуры
    lst_disc = [edu_prog for edu_prog in list_of_eduprogs if edu_prog[2] == l];


    # *   ----- block : сортировка по дисциплинам, практикам, факультативам -----    * #

    # * словарь для хранения всех дисциплин и факультетов где ключ будет -> "код дисциплина или факультатив" -> 
    discs_facults = {}

    facultatives = []
    disciplines = []
    practices = []
    # * распределение на факультативы, дисциплины, практики
    for direction in lst_disc:
        # ! пятая колонка в excel файле обезательно должна представлять -> 'код дисциплины' -> пример 2.1.11.1(Ф) или 2.2.2(П)
        # ! 17 колонка должна содержать примечание на указание факультатив это или нет -> пример 'факультатив' или 'выбор' и т.д.
        # добавление в факультативы
        if "ф" in direction[5].lower(): facultatives.append(direction);
        elif len(direction) >= 17 and "факульт" in direction[16].lower(): facultatives.append(direction);
        # добавление в дисциплины
        elif "п" in direction[5].lower(): practices.append(direction);
        # если не подходит для практики или факультатива тогда в дисциплину
        else: disciplines.append(direction);

        # * добавление ДИСЦИПЛИН и ФАКУЛЬТАТИВОВ для конкретного учебного плана в словарь discs_facults
        # ! в колонках 6 и 7 дожны быть следующие данные : -> 6 - код дисциплины, 7 - наименование дисциплины, 17 - примечание (если факультатив)
        if "ф" in direction[5].lower() or len(direction) >=17 and "факульт" in direction[16].lower():
            discs_facults[(direction[5], direction[6])] = 1
        elif "п" not in direction[5].lower(): discs_facults[(direction[5], direction[6])] = 1

    # *   ----- end : сортировка по дисциплинам, практикам, факультативам -----    * #


    # *   ----- block : создание в word-документе раздела для дисциплин и факультативов -----      * #

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    section = document.sections[0]
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)

    # * добавление параграфа в самое начало документа
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = document.styles['Normal']

    # * направление подготовки и форма обучения
    p.add_run("Направление " + lst_disc[0][3]).bold = True;
    p.add_run("").add_break(); # разделение строки
    p.add_run("очная форма обучения").add_break(); # ! на 22-23 учеб. план только очники

    # * Добавление заголовков Аннотаций
    p.add_run('Аннотации к рабочим программам дисциплин').add_break()
    p.add_run('Аннотации к программам практик').add_break()

    # * добавление второго параграфа
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.add_run('С полнотекстовыми версиями рабочих программ соответствующих дисциплин, программ практик и программы государственной итоговой аттестации можно ознакомиться в ЭБС СибАДИ. ')

    # * добавление третьего параграфа
    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.style = document.styles['Normal']
    p3.add_run('Аннотации к рабочим программам дисциплин').bold = True

    # * итерация по факультативам и дисциплинам
    for disc_or_facult in discs_facults:

        # * временное хранилище всех строк где наименование дисциплины или факультета идентичны с значение disc_or_facult
        tmp_disc_lst = [item for item in lst_disc if item[6] == disc_or_facult[1]]


        # *   ----- block : ЗАГОЛОВОК дисциплины или факультатива -----      * #
        
        # * добавление параграфа дисциплины или факультатива на каждой итерации

        p_disc_facult_title = document.add_paragraph();
        p_disc_facult_title.alignment = WD_ALIGN_PARAGRAPH.CENTER;
        p_disc_facult_title.style = document.styles["Normal"];
        p_disc_facult_title.add_run(f"Дисциплина \" {disc_or_facult[1]}\"")

        # *   ----- end : ЗАГОЛОВОК дисциплины или факультатива -----      * #


        # *   ----- block : ТЕЛО дисциплины или факультатива -----      * #

        # * добавление параграфа ТЕЛА раздела дисциплины или факультатива 
        p_disc_facult_p1 = document.add_paragraph();
        p_disc_facult_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_disc_facult_p1.style = document.styles["Normal"];

        # * определение начала ТЕЛА параграфа дисциплины или факультатива : дисциплина, выборочные дисциплины, факультатив
        if len(tmp_disc_lst[0]) > 16 and ("выбор" in tmp_disc_lst[0][16].lower()):
            p_disc_facult_p1.add_run("Дисциплина относится к вариативной части Блока 1 «Дисциплины (модули). ");
        elif len(tmp_disc_lst[0]) > 16 and ("факультатив" in tmp_disc_lst[0][16].lower()):
            p_disc_facult_p1.add_run("Дисциплина относится к факультативным дисциплинам Блока «ФТД. Факультативы». ");
        else:
            p_disc_facult_p1.add_run('Дисциплина относится к обязательной части Блока 1 «Дисциплины (модули)». ')


            # * ---- start: часть в теле раздела описывающая трудоемкость

        # ! трудоемкость в 10 колонке excel-файла
        complexity: int = int(tmp_disc_lst[0][9]); # convert to integer

        p_disc_facult_p1.add_run(f"Общая трудоемкость освоения дисциплины составляет {complexity}");

        if complexity == 1: p_disc_facult_p1.add_run(" зачетную единицу. ");
        elif complexity == 2: p_disc_facult_p1.add_run(" зачетные единицы. ");
        else: p_disc_facult_p1.add_run(" зачетных единиц. ");

            # * ---- end: часть в теле раздела описывающая трудоемкость

        p_disc_facult_p1.add_run("Рабочей программой дисциплины предусмотрены следующие виды контроля: текущий контроль успеваемости, ");

        # ! 11 и 12 колонка определяют для дисциплины -> зачет или экзамен
        # ! только в одной из них может быть значение цифра = 1 указывающая что это зачет или практика
        # ! исключение составляет только предметы по практике в этом случаи в колонках 11 и 12 будут цифра = 0
        colomun_exam: int = int(tmp_disc_lst[0][10]);
        colomun_test: int = int(tmp_disc_lst[0][11]);

        # * временное хранилеще для определения семестров где будет конкретная дисциплини или факультатив
        storage_semeters_disc: list[str] = [];

        # * итерация по всем дисциплинам или факультетам где они идентичны для конкретного учебного плана
        for disc in tmp_disc_lst:
            # ! 8 колонка определяет конкретный семестр
            semester = int(disc[7]);
            storage_semeters_disc.append(f"{semester} семестр")


        if colomun_exam == 1 and colomun_test == 0:
            p_disc_facult_p1.add_run("промежуточная аттестация в виде экзамена");

        if colomun_exam == 0 and colomun_test == 1:
            p_disc_facult_p1.add_run("промежуточная аттестация в виде зачета")
        
        if len(storage_semeters_disc) > 1 : 
            p_disc_facult_p1.add_run(" (" + ", ".join(storage_semeters_disc) + "). ");
        else:
            p_disc_facult_p1.add_run(". ");

        # * конец раздела дисциплины или факультатива
        p_disc_facult_p1.add_run("В рабочей программе приведено учебно-методическое, информационное обеспечение и материально-техническое обеспечение учебной дисциплины, оценочные и методические материалы.");


        # *   ----- end : ТЕЛО дисциплины или факультатива -----      * #

    # *   ----- end : создание в word-документе раздела для дисциплин и факультативов -----      * #



    # *   ----- block : создание в word-документа разделя для практик -----      * #

    p4 = document.add_paragraph();
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER;
    p4.style = document.styles['Normal'];
    p4.add_run("Аннотации к программам практик").bold = True;

    # * итерация по практикам для определнного учебного плана
    for practice in practices:
        # ! 7 колонка в excel-файле должны представлять наименование практики
        p_practice_title = document.add_paragraph()
        p_practice_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_practice_title.style = document.styles['Normal']
        # * добавление наименование практики в word-файл
        p_practice_title.add_run(practice[6])

        # *   ----- block : ТЕЛО практики -----      * #

        p_practice_p2 = document.add_paragraph();
        p_practice_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY;
        p_practice_p2.style = document.styles["Normal"];

        p_practice_p2.add_run("Практика относится к вариативной части Блока 2 «Практики». ");
        p_practice_p2.add_run('Общая трудоемкость освоения программы практики составляет ')

                # ! трудоемкость в 10 колонке excel-файла
        complexity: int = int(tmp_disc_lst[0][9]); # convert to integer

        if complexity == 1: p_practice_p2.add_run(" зачетную единицу. ");
        elif complexity == 2: p_practice_p2.add_run(" зачетные единицы. ");
        else: p_practice_p2.add_run(" зачетных единиц. ");



        # *   ----- end : ТЕЛО практики -----       * #


        print(practice);

    print("--------------------")

    # *   ----- end : создание в word-документа разделя для практик -----      * #


    
    # *   ----- block : сохранение документа word-аннотации -----      * #

    code_edu_program = lst_disc[0][3].split(" ")[0].replace(".", "")
    name_direction_program = " ".join(lst_disc[0][3].split(" ")[1:]);


    # Создание дополнительной директории по направлению подготовки в директории Аспирантов -> А\Аннотации к РПД_<КодНаправленияПодготовки>_ Научная направленность <Имя направление>
    dirname = 'А\\Аннотации к РПД' + "_" + code_edu_program + "_ Научная направленность " + name_direction_program;
    os.makedirs(dirname, exist_ok=True)


    # имя файла для сохранения в формате .docx в определенной директории А\Аннотации к РПД_<КодНаправленияПодготовки>_ Научная направленность <Имя направление>\<Имя_файла>
    fname = lst_disc[0][2].split('-')[0] + '_Аннотация к РПД_' + lst_disc[0][2].split('-')[1][0:2] +'.docx'
    document.save(dirname+'\\'+fname)

    # *     ----- end : сохранение документа word-аннотации -----     * #



