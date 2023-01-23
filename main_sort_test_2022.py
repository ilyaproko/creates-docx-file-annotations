# This Python file uses the following encoding: utf-8
import xlrd, xlwt
import datetime
import docx
import re
import os
from docx import Document
from docx.shared import Inches, Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

book = xlrd.open_workbook("settings.xls")
sheet = book.sheet_by_index(0)

os.makedirs('БСМ', exist_ok=True)
os.makedirs('А', exist_ok=True)

list_of_eduprogs = []

dict_of_plans = {}

for row_index in range(1, sheet.nrows): #sheet.nrows
    p1, p2, p3, p4, p5, p6, p7, p8, p9, p10, p11, p12, p13, p14, p15 = sheet.row_values(row_index, end_colx=15)
    list_of_eduprogs.append([p2, p3, p4, p5 if p5 != '' else 0 , p6, p7, p8, p9, p10, p11, p12, p13, p14, p15])

#print(len(list_of_eduprogs))

for plan in list_of_eduprogs:
    dict_of_plans[plan[0]] = 1


for l in dict_of_plans:
    lst_disc = [item for item in list_of_eduprogs if item[0] == l]
    #print(l)
    gia_b3 = []
    for item in [item for item in lst_disc if item[1].startswith('Б3.01')]:#Б3.01(Д)
        gia_b3.append(item)
    #print(gia_b3)
    single_disc = {}
    for item in [item for item in lst_disc if (item[1].startswith('Б1') or item[1].startswith('ФТД')) and item[5] != '0.0']: #
        single_disc[(item[1], item[2])] = 1

#####Sorting
    single_disc_sorted = []

    disc_sort_b1b = []
    for item in [item for item in single_disc if item[0].startswith('Б1.Б.') or item[0].startswith('Б1.О.') ]:
        if len(item[0].split('.')) == 3:
            disc_sort_b1b.append( (int(item[0].split('.')[2]), item[1], 0) )
        if  len(item[0].split('.')) == 4:
            disc_sort_b1b.append( (int(item[0].split('.')[2]), item[1], int(item[0].split('.')[3])) )
        #disc_sort_b1b.append(item)
    #disc_sort_b1b.sort(key=lambda items: float(items[0].split('Б1.Б.')[1]))
    disc_sort_b1b.sort(key=lambda items: (items[0], items[2]) )
    single_disc_sorted.extend(disc_sort_b1b)

    if lst_disc[0][13] == 'new':
        disc_sort_b1v = []
        disc_sort_b1v_dv = []
        #for item in [item for item in single_disc if item[0].startswith('Б1.В.') and not item[0].startswith('Б1.В.ДВ.')]:
        #    disc_sort_b1v.append(item)

        for item in [item for item in single_disc if item[0].startswith('Б1.В.') and '.ДВ.' not in item[0]]:
            disc_sort_b1v.append(item)  
        disc_sort_b1v.sort(key=lambda items: float(items[0].split('Б1.В.')[1]))
        single_disc_sorted.extend(disc_sort_b1v)

        for item in [item for item in single_disc if item[0].startswith('Б1.В.0') and '.ДВ.' in item[0]]:
            disc_sort_b1v_dv.append(item)
        disc_sort_b1v_dv.sort(key=lambda items: (float(item[0].split('.ДВ.')[0].split('Б1.В.')[1]) , float(item[0].split('.ДВ.')[1]) )  )
        single_disc_sorted.extend(disc_sort_b1v_dv)
        
    
    if lst_disc[0][13] == 'old':
        disc_sort_b1v = []
        for item in [item for item in single_disc if item[0].startswith('Б1.В.') and not item[0].startswith('Б1.В.ДВ.')]:
            disc_sort_b1v.append(item)
        disc_sort_b1v.sort(key=lambda items: float(items[0].split('Б1.В.')[1]))
        single_disc_sorted.extend(disc_sort_b1v)

    disc_sort_b1vod = []
    for item in [item for item in single_disc if item[0].startswith('Б1.В.ОД.')]:
        disc_sort_b1vod.append(item)
    disc_sort_b1vod.sort(key=lambda items: float(items[0].split('Б1.В.ОД.')[1]))
    single_disc_sorted.extend(disc_sort_b1vod)
    #print(lst_disc[0][0][4:5], lst_disc[0][0])
    if ((lst_disc[0][0][3:4] == '3' or lst_disc[0][0][3:4] == '5') and 'z' not in lst_disc[0][0] and 'G' not in lst_disc[0][0]) or lst_disc[0][0].startswith('p'):
        single_disc_sorted.extend([('Б1.ЭФК','ЭФК')])
    if (lst_disc[0][0][4:5] == '3' or lst_disc[0][0][4:5] == '5') and 'z' in lst_disc[0][0] and 'G' not in lst_disc[0][0]:
        single_disc_sorted.extend([('Б1.ЭФК','ЭФКз')])
    
    disc_sort_b1vdv = []
    for item in [item for item in single_disc if item[0].startswith('Б1.В.ДВ.')]:
        disc_sort_b1vdv.append(item)
    disc_sort_b1vdv.sort(key=lambda items: float(items[0].split('Б1.В.ДВ.')[1]))
    single_disc_sorted.extend(disc_sort_b1vdv)
    

    disc_sort_ftd = []
    for item in [item for item in single_disc if item[0].startswith('ФТД')]:
        disc_sort_ftd.append(item)
    disc_sort_ftd.sort(key=lambda items: float(items[0].split('.')[-1]))
    single_disc_sorted.extend(disc_sort_ftd)
    
    single_disc = {}
    for item in single_disc_sorted:
        single_disc[item[1]] = 1
    
    single_prac = {}
    for item in [item for item in lst_disc if item[1].startswith('Б2') and item[5] != '0.0']:
        single_prac[(item[1], item[2])] = 1
    
    single_prac_sorted = []

    prac_sort_b2u = []
    for item in [item for item in single_prac if '(У)' in item[0]]: #Б2.Б.01(У)
        prac_sort_b2u.append(item)
    prac_sort_b2u.sort(key=lambda items: float(items[0].replace('(У)','').split('.')[-1]))
    single_prac_sorted.extend(prac_sort_b2u)

    prac_sort_b2p = []
    for item in [item for item in single_prac if '(П)' in item[0]]:
        prac_sort_b2p.append(item)
    prac_sort_b2p.sort(key=lambda items: float(items[0].replace('(П)','').split('.')[-1]))
    single_prac_sorted.extend(prac_sort_b2p)
    
    prac_sort_b2pd = []
    for item in [item for item in single_prac if '(Пд)' in item[0]]:
        prac_sort_b2pd.append(item)
    prac_sort_b2pd.sort(key=lambda items: float(items[0].replace('(Пд)','').split('.')[-1]))
    single_prac_sorted.extend(prac_sort_b2pd)

    #prac_sort_b2 = []
    #for item in [item for item in single_prac if item[0] == 'Б2']:
    #    prac_sort_b2p.append(item)
    #single_prac_sorted.extend(prac_sort_b2p)
    
    single_prac = {}
    for item in single_prac_sorted:
        single_prac[item[1]] = 1
    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    section = document.sections[0]
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = document.styles['Normal']
    if lst_disc[0][11].lower().rfind('профиль') != -1:
        f_p = lst_disc[0][11].lower().rfind('профиль')
    elif lst_disc[0][11].lower().rfind('специализация') != -1:
        f_p = lst_disc[0][11].lower().rfind('специализация')
    elif lst_disc[0][11].lower().rfind('магистерская программа') != -1:
        f_p = lst_disc[0][11].lower().rfind('магистерская программа')
    elif lst_disc[0][11].lower().rfind('научная направленность') != -1:
        f_p = lst_disc[0][11].lower().rfind('научная направленность')
    #print(f_p)
    p.add_run( lst_disc[0][11][ : f_p ] ).bold = True
    p.add_run('').add_break()
    p.add_run( lst_disc[0][11][ f_p : ] ).bold = True
    p.add_run('').add_break()
    if 'z' in lst_disc[0][0]:
        p.add_run('заочная форма обучения').add_break()
    elif lst_disc[0][0].startswith('p'):
        p.add_run('очно-заочная форма обучения').add_break()
    else:
        p.add_run('очная форма обучения').add_break()
    p.add_run('Аннотации к рабочим программам дисциплин').add_break()
    p.add_run('Аннотации к программам практик').add_break()
    p.add_run('Аннотация к программе государственной итоговой аттестации')
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.add_run('С полнотекстовыми версиями рабочих программ соответствующих дисциплин, программ практик и программы государственной итоговой аттестации можно ознакомиться в ЭБС СибАДИ. ')

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.style = document.styles['Normal']
    p3.add_run('Аннотации к рабочим программам дисциплин').bold = True
    for disc in single_disc:
        tmp_disc_lst = [item for item in lst_disc if item[2] == disc]
        print(len(tmp_disc_lst));
        print("--------------")
        if disc == 'ЭФК':
            pass
        #    lst_fk_disc = [item for item in lst_disc if item[2] == 'Элективные курсы по физической культуре и спорту']
        #    p_disc_title_fk = document.add_paragraph()
        #    p_disc_title_fk.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #    p_disc_title_fk.style = document.styles['Normal']
        #    p_disc_title_fk.add_run('Дисциплина «Элективные курсы по физической культуре и спорту» (Общая физическая подготовка и спецмедгруппа, Настольный теннис, Спортивные игры, Футбол, Силовые виды спорта, Аэробика)')
        #    p_disc_fk_p1 = document.add_paragraph()
        #    p_disc_fk_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #    p_disc_fk_p1.style = document.styles['Normal']
        #    p_disc_fk_p1.add_run('Дисциплина относится к дисциплинам вариативной части Блока 1 «Дисциплины (модули)». ')
        #    p_disc_fk_p1.add_run('Общая трудоемкость освоения дисциплины составляет 328 часов. ')
        #    p_disc_fk_p1.add_run('Рабочей программой дисциплины предусмотрены следующие виды контроля: промежуточная аттестация в виде зачета (')
        #    #p_disc_fk_z_p1.add_run(str(len(lst_fk_disc)) + ' курс')
        #    l = [str(item+1) for item in range(len(lst_fk_disc))]
        #    p_disc_fk_p1.add_run(', '.join(l))
        #    p_disc_fk_p1.add_run(' семестр). ')
        #    p_disc_fk_p1.add_run('В рабочей программе приведено учебно-методическое, информационное обеспечение и материально-техническое обеспечение учебной дисциплины, оценочные и методические материалы.')
            
        elif disc == 'ЭФКз':
            lst_fk_disc = [item for item in lst_disc if item[2] == 'Элективные курсы по физической культуре и спорту']
            
            #print(lst_fk_disc)
            p_disc_title_fk_z = document.add_paragraph()
            p_disc_title_fk_z.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_disc_title_fk_z.style = document.styles['Normal']
            p_disc_title_fk_z.add_run('Дисциплина «Элективные курсы по физической культуре и спорту»')
            p_disc_fk_z_p1 = document.add_paragraph()
            p_disc_fk_z_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_disc_fk_z_p1.style = document.styles['Normal']
            p_disc_fk_z_p1.add_run('Дисциплина относится к дисциплинам вариативной части Блока 1 «Дисциплины (модули)». ')
            p_disc_fk_z_p1.add_run('Общая трудоемкость освоения дисциплины составляет 328 часов. ')
            p_disc_fk_z_p1.add_run('Рабочей программой дисциплины предусмотрены следующие виды контроля: промежуточная аттестация в виде зачета (')
            p_disc_fk_z_p1.add_run(str(len(lst_fk_disc)) + ' курс')
            p_disc_fk_z_p1.add_run('). ')
            p_disc_fk_z_p1.add_run('В рабочей программе приведено учебно-методическое, информационное обеспечение и материально-техническое обеспечение учебной дисциплины, оценочные и методические материалы.')
        else:
            #print(disc)
            p_disc_title = document.add_paragraph()
            p_disc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_disc_title.style = document.styles['Normal']
            p_disc_title.add_run('Дисциплина "')
            p_disc_title.add_run(tmp_disc_lst[0][2])
            p_disc_title.add_run('"')
            
            p_disc_p1 = document.add_paragraph()
            p_disc_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_disc_p1.style = document.styles['Normal']
            ######old
            
            if tmp_disc_lst[0][13] == 'old':
                if '.Б.' in tmp_disc_lst[0][1]:
                    p_disc_p1.add_run('Дисциплина относится к базовой части Блока 1 «Дисциплины (модули)». ')
                elif '.В.ДВ.' in tmp_disc_lst[0][1]:
                    p_disc_p1.add_run('Дисциплина относится к дисциплинам по выбору вариативной части Блока 1 «Дисциплины (модули)». ')
                elif '.В.' in tmp_disc_lst[0][1] and '.В.ДВ.' not in tmp_disc_lst[0][1]:
                    p_disc_p1.add_run('Дисциплина относится к вариативной части Блока 1 «Дисциплины (модули). ')
                elif 'ФТД' in tmp_disc_lst[0][1]:
                    p_disc_p1.add_run('Дисциплина относится к факультативным дисциплинам Блока «ФТД. Факультативы». ')

            ######new
            if tmp_disc_lst[0][13] == 'new':
                if 'Б1.О.' in tmp_disc_lst[0][1]:
                    p_disc_p1.add_run('Дисциплина относится к обязательной части Блока 1 «Дисциплины (модули)». ')
                elif 'Б1.В.' in tmp_disc_lst[0][1]:
                    p_disc_p1.add_run('Дисциплина относится к части, формируемой участниками образовательных отношений,  Блока 1 «Дисциплины (модули)». ')
                elif 'ФТД' in tmp_disc_lst[0][1]:
                    p_disc_p1.add_run('Дисциплина относится к факультативным дисциплинам Блока «ФТД. Факультативы». ')

            
            p_disc_p1.add_run('Общая трудоемкость освоения дисциплины составляет ')
            #print(tmp_disc_lst[0][5])
            p_disc_p1.add_run(str(tmp_disc_lst[0][5]).replace('.0',''))
            #print(tmp_disc_lst[0][2])
            if tmp_disc_lst[0][2].startswith('Элективные курсы по физической культуре'):
                p_disc_p1.add_run(' часов. ')
            else:
                if str(tmp_disc_lst[0][5]).replace('.0','').endswith('1'):
                    p_disc_p1.add_run(' зачетную единицу. ')
                elif str(tmp_disc_lst[0][5]).replace('.0','').endswith('2') or str(tmp_disc_lst[0][5]).replace('.0','').endswith('3') or str(tmp_disc_lst[0][5]).replace('.0','').endswith('4'):
                    p_disc_p1.add_run(' зачетные единицы. ')
                else:
                    p_disc_p1.add_run(' зачетных единиц. ')
            if 'z' not in tmp_disc_lst[0][0]:
                p_disc_p1.add_run('Рабочей программой дисциплины предусмотрены следующие виды контроля: текущий контроль успеваемости, ')
            else:
                p_disc_p1.add_run('Рабочей программой дисциплины предусмотрены следующие виды контроля: ')
            disc_z_s = [] # 7
            disc_ex_s = [] # 6
            disc_kr_s = [] # 9
            disc_kp_s = [] # 8
            disc_z_w_mark_s = [] #10
            if 'z' not in lst_disc[0][0]: 
                for i in tmp_disc_lst:
                    if i[6] != 0:
                        disc_ex_s.append(str(i[3]) + ' семестр')
                    if i[7] != 0:
                        disc_z_s.append(str(i[3]) + ' семестр')
                    if i[8] != 0:
                        disc_kp_s.append(str(i[3]) + ' семестр')
                    if i[9] != 0:
                        disc_kr_s.append(str(i[3]) + ' семестр')
                    if i[10] != 0:
                        disc_z_w_mark_s.append(str(i[3]) + ' семестр')
                #if tmp_disc_lst[0][2] == 'Композиционное моделирование':
                #    print(disc_ex_s, disc_z_s, disc_kp_s, disc_kr_s, disc_z_w_mark_s)
                    
            else:
                for i in tmp_disc_lst:
                    if i[6] != 0:
                        disc_ex_s.append(str(i[4]) + ' курс')
                    if i[7] != 0:
                        disc_z_s.append(str(i[4]) + ' курс')
                    if i[8] != 0:
                        disc_kp_s.append(str(i[4]) + ' курс')
                    if i[9] != 0:
                        disc_kr_s.append(str(i[4]) + ' курс')
                    if i[10] != 0:
                        disc_z_w_mark_s.append(str(i[4]) + ' курс')
            
            if len(disc_z_s) == 1 and len(disc_ex_s) == 0 and len(disc_kr_s) == 0 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета. ')
            if len(disc_z_s) > 1 and len(disc_ex_s) == 0 and len(disc_kr_s) == 0 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета (')
                p_disc_p1.add_run(', '.join(disc_z_s))
                p_disc_p1.add_run('). ')

            if len(disc_z_s) == 0 and len(disc_ex_s) == 1 and len(disc_kr_s) == 0 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде экзамена. ')
            if len(disc_z_s) ==0 and len(disc_ex_s) > 1 and len(disc_kr_s) == 0 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run('). ')
            
            if len(disc_z_s) == 1 and len(disc_ex_s) == 0 and len(disc_kr_s) == 1 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета ')
                p_disc_p1.add_run('(' + disc_z_s[0] + ') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсовой работы (' + disc_kr_s[0] + '). ')
                
            if len(disc_z_s) == 1 and len(disc_ex_s) == 0 and len(disc_kr_s) == 0 and len(disc_kp_s) == 1 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета ')
                p_disc_p1.add_run('(' + disc_z_s[0] + ') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсового проекта (' + disc_kp_s[0] + '). ')
                
            if len(disc_z_s) > 1 and len(disc_ex_s) == 0 and len(disc_kr_s) == 1 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета (')
                p_disc_p1.add_run(', '.join(disc_z_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсовой работы (' + disc_kr_s[0] + '). ')

            if len(disc_z_s) > 1 and len(disc_ex_s) == 0 and len(disc_kr_s) == 0 and len(disc_kp_s) == 1 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета (')
                p_disc_p1.add_run(', '.join(disc_z_s))
                p_disc_p1.add_run(')')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсового проекта (' + disc_kp_s[0] + '). ')
    #######################################################################################################
            if len(disc_z_s) == 0 and len(disc_ex_s) == 1 and len(disc_kr_s) == 1 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде экзамена ')
                p_disc_p1.add_run('(' + disc_ex_s[0] + ') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсовой работы (' + disc_kr_s[0] + '). ')
                
            if len(disc_z_s) == 0 and len(disc_ex_s) == 1 and len(disc_kr_s) == 0 and len(disc_kp_s) == 1 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде экзамена ')
                p_disc_p1.add_run('(' + disc_ex_s[0] + ') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсового проекта (' + disc_kp_s[0] + '). ')
                
            if len(disc_z_s) == 0 and len(disc_ex_s) > 1 and len(disc_kr_s) == 1 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсовой работы (' + disc_kr_s[0] + '). ')

            if len(disc_z_s) == 0 and len(disc_ex_s) > 1 and len(disc_kr_s) == 0 and len(disc_kp_s) == 1 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсового проекта (' + disc_kp_s[0] + '). ') 
    #######################################################################################################
            if len(disc_z_s) > 0 and len(disc_ex_s) > 0 and len(disc_kr_s) == 0 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета (')
                p_disc_p1.add_run(', '.join(disc_z_s))
                p_disc_p1.add_run(') и ')
                p_disc_p1.add_run('экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run('). ')
                
            if len(disc_z_s) > 0 and len(disc_ex_s) > 0 and len(disc_kr_s) == 1 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета (')
                p_disc_p1.add_run(', '.join(disc_z_s))
                p_disc_p1.add_run('), ')
                p_disc_p1.add_run('экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсовой работы (' + disc_kr_s[0] + '). ')
                
            if len(disc_z_s) > 0 and len(disc_ex_s) > 0 and len(disc_kr_s) == 0 and len(disc_kp_s) == 1 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета (')
                p_disc_p1.add_run(', '.join(disc_z_s))
                p_disc_p1.add_run('), ')
                p_disc_p1.add_run('экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсового проекта (' + disc_kp_s[0] + '). ')

            if len(disc_z_s) > 0 and len(disc_ex_s) > 0 and len(disc_kr_s) > 1 and len(disc_kp_s) > 1 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета (')
                p_disc_p1.add_run(', '.join(disc_z_s))
                p_disc_p1.add_run('), ')
                p_disc_p1.add_run('экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run('), ')
                p_disc_p1.add_run('курсовой работы (')
                p_disc_p1.add_run(', '.join(disc_kr_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсового проекта (')
                p_disc_p1.add_run(', '.join(disc_kp_s))
                p_disc_p1.add_run('). ')
    ############################################### new ###################################################
            if len(disc_z_s) == 0 and len(disc_ex_s) > 0 and len(disc_kr_s) > 0 and len(disc_kp_s) > 0 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run('), ')
                p_disc_p1.add_run('курсовой работы (')
                p_disc_p1.add_run(', '.join(disc_kr_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсового проекта (')
                p_disc_p1.add_run(', '.join(disc_kp_s))
                p_disc_p1.add_run('). ')
    #######################################################################################################
            if len(disc_z_s) == 0 and len(disc_ex_s) > 1 and len(disc_kr_s) == 0 and len(disc_kp_s) > 1 and len(disc_z_w_mark_s) == 0:
                p_disc_p1.add_run('промежуточная аттестация в виде экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('курсового проекта (')
                p_disc_p1.add_run(', '.join(disc_kp_s))
                p_disc_p1.add_run('). ')
    #######################################################################################################
            if len(disc_z_s) == 0 and len(disc_ex_s) > 0 and len(disc_z_w_mark_s) > 0:
                p_disc_p1.add_run('промежуточная аттестация в виде экзамена (')
                p_disc_p1.add_run(', '.join(disc_ex_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('зачета с оценкой (')
                p_disc_p1.add_run(', '.join(disc_z_w_mark_s))
                p_disc_p1.add_run('). ')
    #######################################################################################################
            if len(disc_z_s) == 0 and len(disc_ex_s) == 0 and len(disc_kr_s) == 0 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) > 1:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета с оценкой (')
                p_disc_p1.add_run(', '.join(disc_z_w_mark_s))
                p_disc_p1.add_run('). ')
            if len(disc_z_s) == 0 and len(disc_ex_s) == 0 and len(disc_kr_s) == 0 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) == 1:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета с оценкой. ')
    #######################################################################################################
            if len(disc_z_s) > 1 and len(disc_ex_s) == 0 and len(disc_kr_s) > 1 and len(disc_kp_s) == 0 and len(disc_z_w_mark_s) > 1:
                p_disc_p1.add_run('промежуточная аттестация в виде зачета (')
                p_disc_p1.add_run(', '.join(disc_z_s))
                p_disc_p1.add_run('), ')
                p_disc_p1.add_run('курсовой работы (')
                p_disc_p1.add_run(', '.join(disc_kr_s))
                p_disc_p1.add_run(') ')
                p_disc_p1.add_run('и ')
                p_disc_p1.add_run('промежуточная аттестация в виде зачета с оценкой (')
                p_disc_p1.add_run(', '.join(disc_z_w_mark_s))
                p_disc_p1.add_run('). ')
            
            
    #######################################################################################################
            p_disc_p1.add_run('В рабочей программе приведено учебно-методическое, информационное обеспечение и материально-техническое обеспечение учебной дисциплины, оценочные и методические материалы.')
#!!!!!!!!!!! FK
#####################
    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.style = document.styles['Normal']
    p4.add_run('Аннотации к программам практик').bold = True
    if lst_disc[0][0][3:4] == '6' or (lst_disc[0][0].startswith('z') and lst_disc[0][0][4:5] == '6'):
        p_disc_title = document.add_paragraph()
        p_disc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_disc_title.style = document.styles['Normal']
        p_disc_title.add_run('Педагогическая практика')

        p_disc_p1 = document.add_paragraph()
        p_disc_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_disc_p1.style = document.styles['Normal']
        p_disc_p1.add_run('Практика относится к вариативной части Блока 2 «Практики». Общая трудоемкость освоения программы практики составляет 6 зачетных единиц. Рабочей программой практики предусмотрены следующие виды контроля: промежуточная аттестация в виде зачета с оценкой. В программе практики приведено учебно-методическое, информационное обеспечение и материально-техническое обеспечение практики, оценочные и методические материалы.')

        p_disc_title2 = document.add_paragraph()
        p_disc_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_disc_title2.style = document.styles['Normal']
        p_disc_title2.add_run('Практика по получению профессиональных умений и опыта профессиональной деятельности')

        p_disc_p2 = document.add_paragraph()
        p_disc_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_disc_p2.style = document.styles['Normal']
        p_disc_p2.add_run('Практика относится к вариативной части Блока 2 «Практики». Общая трудоемкость освоения программы практики составляет 4 зачетные единицы. Рабочей программой практики предусмотрены следующие виды контроля: промежуточная аттестация в виде зачета с оценкой. В программе практики приведено учебно-методическое, информационное обеспечение и материально-техническое обеспечение практики, оценочные и методические материалы.')
    for prac in single_prac:
        tmp_prac_lst = [item for item in lst_disc if item[2] == prac]
        p_disc_title = document.add_paragraph()
        p_disc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_disc_title.style = document.styles['Normal']
        p_disc_title.add_run(tmp_prac_lst[0][2])

        p_disc_p2 = document.add_paragraph()
        p_disc_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_disc_p2.style = document.styles['Normal']
        ###old
        if tmp_prac_lst[0][13] == 'old':
            if tmp_prac_lst[0][12] == 'В':
                p_disc_p2.add_run('Практика относится к вариативной части Блока 2 «Практики». ')
            else:
                p_disc_p2.add_run('Практика относится к базовой части Блока 2 «Практики». ')

        ###new
        if tmp_prac_lst[0][13] == 'new':
            if tmp_prac_lst[0][1].startswith('Б2.О.'):
                p_disc_p2.add_run('Практика относится к обязательной части Блока 2 «Практики». ')
            else:
                p_disc_p2.add_run('Практика относится к части, формируемой участниками образовательных отношений,  Блока 2 «Практики». ')
        
        p_disc_p2.add_run('Общая трудоемкость освоения программы практики составляет ')
        p_disc_p2.add_run(str(tmp_prac_lst[0][5]).replace('.0',''))

        if int(str(tmp_prac_lst[0][5]).replace('.0','')) < 10:
            if str(tmp_prac_lst[0][5]).replace('.0','').endswith('1'):
                p_disc_p2.add_run(' зачетную единицу. ')
            elif str(tmp_prac_lst[0][5]).replace('.0','').endswith('2') or str(tmp_prac_lst[0][5]).replace('.0','').endswith('3') or str(tmp_prac_lst[0][5]).replace('.0','').endswith('4'):
                p_disc_p2.add_run(' зачетные единицы. ')
            else:
                p_disc_p2.add_run(' зачетных единиц. ')
        else:
            p_disc_p2.add_run(' зачетных единиц. ')
        p_disc_p2.add_run('Рабочей программой практики предусмотрены следующие виды контроля: промежуточная аттестация в виде зачета с оценкой. В программе практики приведено учебно-методическое, информационное обеспечение и материально-техническое обеспечение практики, оценочные и методические материалы.')
    
    if (lst_disc[0][0][3:4] == '6' or (lst_disc[0][0].startswith('z') and lst_disc[0][0][4:5] == '6')):
        p5 = document.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.style = document.styles['Normal']
        p5.add_run('Аннотация к программе научных исследований').bold = True
        p_gia_title = document.add_paragraph()
        p_gia_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_gia_title.style = document.styles['Normal']
        p_gia_title.add_run('Научно-исследовательская деятельность и подготовка научно-квалификационной работы (диссертации)')
        p_disc_p3 = document.add_paragraph()
        p_disc_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_disc_p3.style = document.styles['Normal']
        p_disc_p3.add_run('Научно-исследовательская деятельность и подготовка научно-квалификационной работы (диссертации) относится к вариативной части Блока 3 «Научные исследования». Она реализуется с первого курса (в соответствии с учебным планом) обучения,  параллельно с изучением всех дисциплин учебного плана. Общая трудоемкость освоения составляет 191 зачетную единицу. Рабочей программой научных исследований предусмотрены следующие виды контроля: промежуточная аттестация в виде зачета с оценкой (2, 3, 4, 5, 6, 7 семестры). В рабочей программе приведено учебно-методическое, информационное обеспечение и материально-техническое обеспечение научных исследований, оценочные и методические материалы.')

        p6 = document.add_paragraph()
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p6.style = document.styles['Normal']
        p6.add_run('Аннотация к программе ГИА').bold = True
        p_gia_title2 = document.add_paragraph()
        p_gia_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_gia_title2.style = document.styles['Normal']
        p_gia_title2.add_run('Государственная итоговая аттестация')
        p_disc_p4 = document.add_paragraph()
        p_disc_p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_disc_p4.style = document.styles['Normal']
        p_disc_p4.add_run('Государственная итоговая аттестация, включая подготовку к сдаче и сдачу государственного экзамена, а также представление научного доклада об основных результатах подготовленной научно-квалификационной работы (диссертации), относится к базовой части Блока 4 «Государственная итоговая аттестация». К государственной итоговой аттестации допускается обучающийся, не имеющий академической задолженности и в полном объеме выполнивший учебный план или индивидуальный учебный план по соответствующей образовательной программе высшего образования. Программой государственной итоговой аттестации предусмотрена форма проведения ГИА и вид научного доклада об основных результатах подготовленной научно-квалификационной работы (диссертации). В программе приведены примерный перечень направлений научных исследований, оценочные и методические материалы.')
    else:
        p5 = document.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.style = document.styles['Normal']
        p5.add_run('Аннотация к программе ГИА').bold = True
        p_gia_title = document.add_paragraph()
        p_gia_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_gia_title.style = document.styles['Normal']
        p_gia_title.add_run('Государственная итоговая аттестация')
        p_disc_p3 = document.add_paragraph()
        p_disc_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_disc_p3.style = document.styles['Normal']
        #print(gia_b3[0][2])
        p_disc_p3.add_run(gia_b3[0][2])
        p_disc_p3.add_run(' относится к базовой части Блока 3 «Государственная итоговая аттестация». К государственной итоговой аттестации допускается обучающийся, не имеющий академической задолженности и в полном объеме выполнивший учебный план или индивидуальный учебный план по соответствующей образовательной программе высшего образования. Программой государственной итоговой аттестации предусмотрена форма проведения ГИА и вид выпускной квалификационной работы. В программе приведен макет выпускной квалификационной работы, оценочные и методические материалы.')
        #p_disc_p3.add_run('Защита выпускной квалификационной работы, включая подготовку к процедуре защиты и процедуру защиты относится к базовой части Блока 3 «Государственная итоговая аттестация». К государственной итоговой аттестации допускается обучающийся, не имеющий академической задолженности и в полном объеме выполнивший учебный план или индивидуальный учебный план по соответствующей образовательной программе высшего образования. Программой государственной итоговой аттестации предусмотрена форма проведения ГИА и вид выпускной квалификационной работы. В программе приведен макет выпускной квалификационной работы, оценочные и методические материалы.')

    # if (lst_disc[0][0][3:4] != '6' ):
    #     p5 = document.add_paragraph()
    #     p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     p5.style = document.styles['Normal']
    #     p5.add_run('Аннотация к программе ГИА').bold = True
    #     p_gia_title = document.add_paragraph()
    #     p_gia_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     p_gia_title.style = document.styles['Normal']
    #     p_gia_title.add_run('Государственная итоговая аттестация')
    #     p_disc_p3 = document.add_paragraph()
    #     p_disc_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #     p_disc_p3.style = document.styles['Normal']
    #     p_disc_p3.add_run('Защита выпускной квалификационной работы, включая подготовку к процедуре защиты и процедуру защиты относится к базовой части Блока 3 «Государственная итоговая аттестация». К государственной итоговой аттестации допускается обучающийся, не имеющий академической задолженности и в полном объеме выполнивший учебный план или индивидуальный учебный план по соответствующей образовательной программе высшего образования. Программой государственной итоговой аттестации предусмотрена форма проведения ГИА и вид выпускной квалификационной работы. В программе приведен макет выпускной квалификационной работы, оценочные и методические материалы.')
    # else:
    #     p5 = document.add_paragraph()
    #     p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     p5.style = document.styles['Normal']
    #     p5.add_run('Аннотация к программе научных исследований').bold = True
    #     p_gia_title = document.add_paragraph()
    #     p_gia_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     p_gia_title.style = document.styles['Normal']
    #     p_gia_title.add_run('Научно-исследовательская деятельность и подготовка научно-квалификационной работы (диссертации)')
    #     p_disc_p3 = document.add_paragraph()
    #     p_disc_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #     p_disc_p3.style = document.styles['Normal']
    #     p_disc_p3.add_run('Научно-исследовательская деятельность и подготовка научно-квалификационной работы (диссертации) относится к вариативной части Блока 3 «Научные исследования». Она реализуется с первого курса (в соответствии с учебным планом) обучения,  параллельно с изучением всех дисциплин учебного плана. Общая трудоемкость освоения составляет 191 зачетную единицу. Рабочей программой научных исследований предусмотрены следующие виды контроля: промежуточная аттестация в виде зачета с оценкой (2, 3, 4, 5, 6, 7 семестры). В рабочей программе приведено учебно-методическое, информационное обеспечение и материально-техническое обеспечение научных исследований, оценочные и методические материалы.')

    #     p6 = document.add_paragraph()
    #     p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     p6.style = document.styles['Normal']
    #     p6.add_run('Аннотация к программе ГИА').bold = True
    #     p_gia_title2 = document.add_paragraph()
    #     p_gia_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     p_gia_title2.style = document.styles['Normal']
    #     p_gia_title2.add_run('Государственная итоговая аттестация')
    #     p_disc_p4 = document.add_paragraph()
    #     p_disc_p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #     p_disc_p4.style = document.styles['Normal']
    #     p_disc_p4.add_run('Государственная итоговая аттестация, включая подготовку к сдаче и сдачу государственного экзамена, а также представление научного доклада об основных результатах подготовленной научно-квалификационной работы (диссертации), относится к базовой части Блока 4 «Государственная итоговая аттестация». К государственной итоговой аттестации допускается обучающийся, не имеющий академической задолженности и в полном объеме выполнивший учебный план или индивидуальный учебный план по соответствующей образовательной программе высшего образования. Программой государственной итоговой аттестации предусмотрена форма проведения ГИА и вид научного доклада об основных результатах подготовленной научно-квалификационной работы (диссертации). В программе приведены примерный перечень направлений научных исследований, оценочные и методические материалы.')
    
    #Наименование папки "Аннотации к РПД_(шифр направления)_(профиль).
    #dirname = 'Аннотации к РПД_' + lst_disc[0][11][ : f_p ].strip() + '_' + lst_disc[0][11][ f_p : ].strip().replace('"','')

    

    if lst_disc[0][0][3:4] == '6' or (lst_disc[0][0].startswith('z') and lst_disc[0][0][4:5] == '6'):
        dirname = 'А\\Аннотации к РПД_' + lst_disc[0][0].split('-')[0] + '_' + lst_disc[0][11][ f_p : ].strip().replace('"','').replace(':','').strip()
        os.makedirs(dirname, exist_ok=True)
    else:
        dirname = 'БСМ\\Аннотации к РПД_' + lst_disc[0][0].split('-')[0] + '_' + lst_disc[0][11][ f_p : ].strip().replace('"','').replace(':','').strip()
        os.makedirs(dirname, exist_ok=True)

    
    if 'z' in lst_disc[0][0]:
        eduform = 'ЗАОЧНАЯ'
    elif lst_disc[0][0].startswith('p'):
        eduform = 'ОЧНО-ЗАОЧНАЯ'
    else:
        eduform = 'ОЧНАЯ'
    #fname = lst_disc[0][0][0:6] + '_Аннотация к РПД_' + eduform + '_' + lst_disc[0][0].split('-')[1][0:2] +'.docx'
    fname = lst_disc[0][0].split('-')[0] + '_Аннотация к РПД_' + eduform + '_' + lst_disc[0][0].split('-')[1][0:2] +'.docx'    
    #document.save(l+'.docx')
    document.save(dirname+'\\'+fname)
    # print(fname)
